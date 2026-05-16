from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor


def create_invoice() -> None:
    document = Document()

    # --- Styles ---
    style = document.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # --- Header Section ---
    # Title
    head = document.add_heading("INVOICE", 0)
    head.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Your Details (Top Left) vs Invoice Meta (Top Right)
    table_header = document.add_table(rows=1, cols=2)
    table_header.autofit = False
    table_header.allow_autofit = False
    table_header.columns[0].width = Inches(4.0)
    table_header.columns[1].width = Inches(2.5)

    # Left Cell: Your Info
    cell_sender = table_header.cell(0, 0)
    p = cell_sender.paragraphs[0]
    p.add_run("YOUR NAME / COMPANY\n").bold = True
    p.add_run("Musterstraße 1\n12345 Berlin\nGermany\n\n")
    p.add_run("Email: you@example.com\nWeb: www.your-website.com")

    # Right Cell: Invoice Meta Data
    cell_meta = table_header.cell(0, 1)
    p = cell_meta.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Invoice No: ").bold = True
    p.add_run("2025-101\n")
    p.add_run("Date: ").bold = True
    p.add_run("17 Dec 2025\n")
    p.add_run("Customer ID: ").bold = True
    p.add_run("US-550")

    document.add_paragraph("\n")  # Spacer

    # --- Recipient Section ---
    document.add_paragraph("Bill To:").bold = True
    document.add_paragraph(
        "US Company LLC\n123 Innovation Drive\nSan Francisco, CA 94103\nUSA"
    )

    document.add_paragraph("\n")  # Spacer

    # --- Line Items Table ---
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    # Header Row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Description"
    hdr_cells[1].text = "Quantity"
    hdr_cells[2].text = "Unit Price"
    hdr_cells[3].text = "Total (USD)"

    # Make header bold
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            paragraph.runs[0].bold = True

    # Row 1: The Service
    row_cells = table.add_row().cells
    row_cells[0].text = "Consultancy Services (Project X)"
    row_cells[1].text = "1"
    row_cells[2].text = "$2,500.00"
    row_cells[3].text = "$2,500.00"

    document.add_paragraph("\n")  # Spacer

    # --- Totals Section ---
    # We use a table to align totals to the right
    total_table = document.add_table(rows=3, cols=2)
    total_table.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Net
    total_table.rows[0].cells[0].text = "Net Amount:"
    total_table.rows[0].cells[1].text = "$2,500.00"

    # VAT
    total_table.rows[1].cells[0].text = "VAT (0%):"
    total_table.rows[1].cells[1].text = "$0.00"

    # Grand Total
    total_table.rows[2].cells[0].text = "Total Amount:"
    total_table.rows[2].cells[1].text = "$2,500.00"

    # Bold the grand total
    total_table.rows[2].cells[1].paragraphs[0].runs[0].bold = True

    document.add_paragraph("\n")

    # --- Legal Notes (CRITICAL) ---
    note = document.add_paragraph()
    note_runner = note.add_run("Tax Note / Steuerhinweis:")
    note_runner.bold = True
    note.add_run(
        "\nService not subject to German VAT according to § 3a (2) UStG. Place of supply is the USA. Tax liability shifts to the recipient."
    )
    note.add_run("\n(Nicht im Inland steuerbare Leistung gemäß § 3a Abs. 2 UStG).")

    document.add_paragraph("\n")

    # --- Footer / Payment Info ---
    footer_p = document.add_paragraph()
    footer_p.add_run("Payment Terms: ").bold = True
    footer_p.add_run(
        "Please transfer the amount within 14 days to the following account:\n\n"
    )

    footer_p.add_run("Account Holder: Your Name\n")
    footer_p.add_run("Bank: Your Bank Name\n")
    footer_p.add_run("IBAN: DEXX XXXX XXXX XXXX XXXX XX\n")
    footer_p.add_run("BIC/SWIFT: XXXXXXXX\n\n")

    # Tax IDs
    footer_p.add_run("Tax Number (Steuernummer): 123/456/789\n")
    footer_p.add_run("VAT ID (USt-IdNr): DE123456789")

    # Save
    file_name = "Invoice_Germany_to_USA.docx"
    document.save(file_name)
    print(f"Invoice saved as '{file_name}'")


if __name__ == "__main__":
    create_invoice()
